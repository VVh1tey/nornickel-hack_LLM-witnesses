"""Экспорт бизнес-отчёта DOCX: цель, топ гипотез с обоснованием, источниками и roadmap."""

from __future__ import annotations

from pathlib import Path

from docx import Document

from hypofactory import config
from hypofactory.schemas import Hypothesis

EXPORTS_DIR = config.DATA_DIR / "sessions" / "exports"


def export_docx(session_id: str, goal: str, hypotheses: list[Hypothesis]) -> Path:
    EXPORTS_DIR.mkdir(parents=True, exist_ok=True)
    path = EXPORTS_DIR / f"{session_id}.docx"

    doc = Document()
    doc.add_heading("Фабрика гипотез — отчёт", level=1)
    doc.add_paragraph(f"Цель: {goal}")
    doc.add_paragraph(f"Сгенерировано гипотез: {len(hypotheses)}")

    for i, hyp in enumerate(hypotheses, start=1):
        doc.add_heading(f"{i}. {hyp.statement}", level=2)
        doc.add_paragraph(f"Механизм влияния: {hyp.mechanism}")
        doc.add_paragraph(f"Ожидаемый эффект: {hyp.expected_effect}")

        scores = (
            f"Новизна: {hyp.novelty} | Реализуемость: {hyp.feasibility} | "
            f"Эффект: {hyp.impact} | Риск: {hyp.risk} | Итоговый скор: {hyp.score}"
        )
        doc.add_paragraph(scores)

        if hyp.already_tried:
            doc.add_paragraph(f"Уже пробовали: {hyp.already_tried}")
        if hyp.critic_verdict:
            doc.add_paragraph(f"Вердикт критика: {hyp.critic_verdict}")

        if hyp.sources:
            doc.add_paragraph("Источники:")
            for src in hyp.sources:
                quote = f" — «{src.quote}»" if src.quote else ""
                doc.add_paragraph(f"  • {src.source}{quote}", style="List Bullet")

        if hyp.roadmap:
            doc.add_paragraph("Дорожная карта проверки:")
            for step in hyp.roadmap:
                extra = []
                if step.resources:
                    extra.append(f"ресурсы: {step.resources}")
                if step.success_criteria:
                    extra.append(f"критерий успеха: {step.success_criteria}")
                suffix = f" ({'; '.join(extra)})" if extra else ""
                doc.add_paragraph(f"  {step.step}{suffix}", style="List Number")

    doc.save(str(path))
    return path
